# cbosa_scraper/docx_newsletter.py
import os
import logging
from datetime import datetime
from typing import List, Dict
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH



class DocxNewsletterGenerator:
    def __init__(self, output_dir: str = "./out"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        self.logger = logging.getLogger(__name__)

    def create(self, analyses: List[Dict], search_params: Dict, stats: Dict) -> str:
        """
        Tworzy plik DOCX z analizami AI i zwraca pełną ścieżkę do pliku.
        """
        current_date = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = os.path.join(self.output_dir, f"CBOSA_AI_Biuletyn_{current_date}.docx")

        doc = Document()

        # Nagłówki
        title = doc.add_heading('Biuletyn Analityczny CBOSA', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_heading('Analiza Orzeczeń Sądów Administracyjnych', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Podsumowanie
        date_str = datetime.now().strftime('%d.%m.%Y')
        summary_para = doc.add_paragraph()
        summary_para.add_run("Data: ").bold = True
        summary_para.add_run(f"{date_str}\n")
        summary_para.add_run("Liczba orzeczeń: ").bold = True
        summary_para.add_run(f"{len(analyses)}")

        doc.add_page_break()

        # Treści analiz
        for i, analysis in enumerate(analyses, 1):
            case_info = analysis.get('case_info', {})
            case_url = case_info.get('url') or analysis.get('case_url', '')
            # Źródło (małą czcionką)
            src_para = doc.add_paragraph()
            run = src_para.add_run(f"Źródło: {str(case_url)[:120]}{'...' if case_url and len(case_url) > 120 else ''}")
            run.italic = True
            run.font.size = Pt(9)

            # Tekst analizy
            analysis_text = analysis.get('analysis', '').strip()
            sections = [s for s in analysis_text.split('\n\n') if s.strip()]
            for section in sections:
                if section.isupper() and section.endswith(':'):
                    doc.add_heading(section.strip(), level=2)
                else:
                    p = doc.add_paragraph(section.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if i < len(analyses):
                doc.add_page_break()

        doc.save(output_path)
        self.logger.info("DOCX newsletter generated: %s", output_path)
        return output_path
