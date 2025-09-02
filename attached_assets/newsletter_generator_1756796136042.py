import os
import logging
from datetime import datetime
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

class NewsletterGenerator:
    def __init__(self):
        """Initialize the newsletter generator."""
        self.logger = logging.getLogger(__name__)
        

    
    def create_docx_newsletter(self, analyses: List[Dict], search_params: Dict, stats: Dict, output_path: str) -> str:
        """
        Generate DOCX newsletter with AI analyses.
        
        Args:
            analyses: List of analysis results
            search_params: Original search parameters
            stats: Analysis statistics
            output_path: Path where to save the DOCX
            
        Returns:
            Path to generated DOCX file
        """
        try:
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading('Biuletyn Analityczny CBOSA', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading('Analiza Orzeczeń Sądów Administracyjnych', level=2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Summary information
            current_date = datetime.now().strftime('%d.%m.%Y')
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"Data: ").bold = True
            summary_para.add_run(f"{current_date}\n")
            summary_para.add_run(f"Liczba orzeczeń: ").bold = True
            summary_para.add_run(f"{len(analyses)}")
            
            # Page break
            doc.add_page_break()
            
            # Add each analysis
            for i, analysis in enumerate(analyses, 1):
                # Analysis title
                filename = analysis.get('filename', f'Orzeczenie_{i}')
                title_heading = doc.add_heading(f"Orzeczenie {i}: {filename}", level=1)
                
                # Source URL
                source_para = doc.add_paragraph()
                source_run = source_para.add_run(f"Źródło: {analysis.get('case_url', '')[:80]}...")
                source_run.italic = True
                source_run.font.size = Pt(9)
                
                # Analysis content
                analysis_text = analysis['analysis']
                
                # Split analysis into sections for better formatting
                sections = analysis_text.split('\n\n')
                for section in sections:
                    if section.strip():
                        # Check if section is a heading (starts with caps and ends with colon)
                        if section.strip().isupper() and section.strip().endswith(':'):
                            doc.add_heading(section.strip(), level=2)
                        else:
                            para = doc.add_paragraph(section.strip())
                            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Add page break between analyses (except for the last one)
                if i < len(analyses):
                    doc.add_page_break()
            

            
            # Save document
            doc.save(output_path)
            
            self.logger.info(f"DOCX newsletter generated: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error generating DOCX newsletter: {e}")
            raise
    


    def generate_newsletter(self, format_type: str, analyses: List[Dict], search_params: Dict, stats: Dict, output_dir: str) -> str:
        """
        Generate newsletter in DOCX format.
        
        Args:
            format_type: Ignored - only DOCX format is supported
            analyses: List of analysis results
            search_params: Original search parameters
            stats: Analysis statistics
            output_dir: Directory where to save the file
            
        Returns:
            Path to generated DOCX file
        """
        current_date = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"CBOSA_AI_Biuletyn_{current_date}.docx"
        output_path = os.path.join(output_dir, filename)
        return self.create_docx_newsletter(analyses, search_params, stats, output_path)